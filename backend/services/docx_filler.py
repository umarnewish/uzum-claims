"""Token engine for the Uzum claim/agreement docx templates.

- Scalar tokens: `{{token}}` anywhere in any paragraph (including table
  cells) is replaced. The token's run is collapsed to keep the first
  run's formatting; paragraph alignment / font are preserved.
- Items table: any row whose joined cell text contains `{{item.*}}`
  tokens is treated as the template row. It is cloned once per item
  with `{{item.field}}` resolved against each item dict; the original
  template row is removed.

The engine is deliberately small — it does not implement Jinja loops,
conditionals, or filters. The four real Uzum templates only need scalar
substitution + a single repeating items row.
"""
from __future__ import annotations

import copy
import logging
import re
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.table import Table, _Row

logger = logging.getLogger(__name__)

SCALAR_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
ITEM_RE = re.compile(r"\{\{\s*item\.([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def fill(template_path: str | Path, output_path: str | Path,
         ctx: Mapping[str, Any], items: list[Mapping[str, Any]] | None = None) -> Path:
    """Render `template_path` with `ctx` and `items`, write to `output_path`."""
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # Tables first — items expansion changes paragraph contents that we'll
    # then scalar-substitute on the next pass.
    if items is not None:
        for tbl in doc.tables:
            _expand_items(tbl, items)

    _replace_scalars_in_paragraphs(doc.paragraphs, ctx)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _replace_scalars_in_paragraphs(cell.paragraphs, ctx)

    doc.save(str(output_path))
    return output_path


def _replace_scalars_in_paragraphs(paragraphs, ctx: Mapping[str, Any]) -> None:
    for p in paragraphs:
        joined = "".join(r.text for r in p.runs)
        if "{{" not in joined:
            continue
        new = SCALAR_RE.sub(lambda m: _str(ctx.get(m.group(1), m.group(0))), joined)
        if new != joined:
            _collapse_runs(p, new)


def _expand_items(tbl: Table, items: list[Mapping[str, Any]]) -> None:
    template_row_idx = None
    for ri, row in enumerate(tbl.rows):
        joined = " ".join(c.text for c in row.cells)
        if ITEM_RE.search(joined):
            template_row_idx = ri
            break
    if template_row_idx is None:
        return
    template_row = tbl.rows[template_row_idx]
    template_xml = template_row._tr
    parent = template_xml.getparent()
    insert_at = list(parent).index(template_xml)

    for offset, item in enumerate(items):
        new_tr = copy.deepcopy(template_xml)
        parent.insert(insert_at + offset, new_tr)
        new_row = _Row(new_tr, tbl)
        for cell in new_row.cells:
            for p in cell.paragraphs:
                joined = "".join(r.text for r in p.runs)
                if "{{" not in joined:
                    continue
                resolved = ITEM_RE.sub(
                    lambda m: _str(item.get(m.group(1), m.group(0))), joined
                )
                if resolved != joined:
                    _collapse_runs(p, resolved)

    parent.remove(template_xml)


def _collapse_runs(paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _str(v: Any) -> str:
    if v is None:
        return ""
    return str(v)
