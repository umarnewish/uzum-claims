"""One-shot script to convert the four real Uzum docx templates into
placeholder templates. Idempotent — safe to re-run.

For each template, we find paragraphs containing known seller / claim
data and replace the data with `{{tokens}}`. Run-level formatting is
collapsed onto the first run of each affected paragraph (keeps paragraph
alignment and base font; loses mid-paragraph bold/italic that bracketed
data, which is acceptable here).

The items table in claim_ru.docx keeps its header (row 0) plus one
template row of `{{item.*}}` tokens; remaining data rows are deleted.
The other claim/agreement docs get scalar substitutions only.

Usage: python scripts/reauthor_templates.py
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.table import Table

ROOT = Path(__file__).resolve().parent.parent
TEMPLATES = ROOT / "templates"


def collapse_runs(paragraph, new_text: str) -> None:
    """Replace paragraph contents with `new_text` in a single run that
    inherits the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    first = runs[0]
    first.text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def replace_in_paragraph(paragraph, replacements: list[tuple[str, str]]) -> bool:
    """Try to substitute each (needle, replacement) pair in the joined
    text of the paragraph. Returns True if any substitution happened."""
    joined = "".join(r.text for r in paragraph.runs)
    new = joined
    for needle, repl in replacements:
        if needle in new:
            new = new.replace(needle, repl)
    if new != joined:
        collapse_runs(paragraph, new)
        return True
    return False


def replace_anywhere(doc, replacements: list[tuple[str, str]]) -> int:
    """Walk paragraphs + table cells, applying replacements."""
    count = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, replacements):
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, replacements):
                        count += 1
    return count


def reauthor_claim_ru() -> None:
    src = TEMPLATES / "claim_ru.docx"
    print(f"\n→ {src.name}")
    doc = Document(src)

    # Scalar replacements. Order matters: replace longer strings first so
    # we don't partially replace a substring that appears in another match.
    repls = [
        ("Абдурахимов Фозил Тошкулович", "{{fio}}"),
        ("ALFA POLIMER LINE", "{{legal_name}}"),
        ("ООО «{{legal_name}}»", "{{legal_form}} «{{legal_name}}»"),
        ("306338752", "{{inn}}"),
        ("00083", "{{mfo}}"),
        ("АК Хамкорбанк", "{{bank_name}}"),
        ("20208000305061552001", "{{bank_account}}"),
        ("3 434 732,90", "{{total_amount}}"),
    ]
    n = replace_anywhere(doc, repls)
    print(f"  scalar substitutions: {n}")

    # Items table: keep header (row 0), turn row 1 into the {{item.*}}
    # template row, delete rows 2..N-1.
    if not doc.tables:
        sys.exit("expected items table in claim_ru.docx")
    tbl: Table = doc.tables[0]
    if len(tbl.rows) < 2:
        sys.exit("items table needs at least 2 rows (header + 1 data)")

    template_cells = [
        "{{item.product_title}}",
        "{{item.barcode}}",
        "{{item.reason}}",
        "{{item.unit_compensation}}",
        "{{item.expected_qty}}",
        "{{item.line_total}}",
    ]
    template_row = tbl.rows[1]
    for cell, tok in zip(template_row.cells, template_cells):
        # cell may have multiple paragraphs; reduce to one with the token
        first_p = cell.paragraphs[0]
        # Remove subsequent paragraphs in the cell
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
        collapse_runs(first_p, tok)

    # Delete everything after row 1
    rows_to_remove = list(tbl.rows)[2:]
    for row in rows_to_remove:
        row._element.getparent().remove(row._element)
    print(f"  items table: header + 1 template row (removed {len(rows_to_remove)} data rows)")

    doc.save(src)
    print(f"  saved {src}")


def reauthor_claim_uz() -> None:
    src = TEMPLATES / "claim_uz.docx"
    print(f"\n→ {src.name}")
    doc = Document(src)

    # Blank template — replaces "______..." underscores after each label.
    # Match label + underscores explicitly to avoid clobbering other lines.
    repls = [
        ("KIMDAN (FISh): ______________________________", "KIMDAN (FISh): {{fio}}"),
        ("MChJ/XK: ______________________________", "{{legal_form}}: {{legal_name}}"),
        ("STIR/JShShIR: ______________________________", "STIR/JShShIR: {{inn}}"),
        ("HISOB RAQAMI: ______________________________", "HISOB RAQAMI: {{bank_account}}"),
        ("MFO: ______________________________", "MFO: {{mfo}}"),
        ("_________________________ so‘m", "{{total_amount}} so‘m"),
    ]
    n = replace_anywhere(doc, repls)
    print(f"  scalar substitutions: {n}")

    # Items table — same shape as RU. 6 empty data rows after header.
    if doc.tables:
        tbl = doc.tables[0]
        if len(tbl.rows) >= 2:
            template_cells = [
                "{{item.product_title}}",
                "{{item.barcode}}",
                "{{item.reason}}",
                "{{item.unit_compensation}}",
                "{{item.expected_qty}}",
                "{{item.line_total}}",
            ]
            template_row = tbl.rows[1]
            for cell, tok in zip(template_row.cells, template_cells):
                first_p = cell.paragraphs[0]
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
                collapse_runs(first_p, tok)
            rows_to_remove = list(tbl.rows)[2:]
            for row in rows_to_remove:
                row._element.getparent().remove(row._element)
            print(f"  items table tokenized (removed {len(rows_to_remove)} extra rows)")

    doc.save(src)
    print(f"  saved {src}")


def reauthor_agreement_ru() -> None:
    """Agreement RU has scalar PII split across multiple <w:t> runs that
    the python-docx joined-paragraph replacer misses. Use raw XML
    replacement on the concatenated paragraph text and collapse runs."""
    src = TEMPLATES / "agreement_ru.docx"
    print(f"\n→ {src.name}")

    repls = [
        # Order matters — longest string first so partial substrings don't
        # pre-bind. The {{token}} placeholders match the keys docx_filler
        # substitutes from the seller_profile + claim row.
        ("Дополнительное соглашение №", "Дополнительное соглашение №{{claim_no}}"),
        ("Абдурахимов Фозил Тошкулович", "{{fio}}"),
        ("г.Ташкент. Чиланзарский район, Гулистан 6-15", "{{address}}"),
        ("Чиланзарский район, Гулистан 6-15", "{{address}}"),
        ("ИП___ООО «ALFA POLIMER LINE»", "{{legal_form}} «{{legal_name}}»"),
        ('"ALFA POLIMER LINE"', "«{{legal_name}}»"),
        ("ALFA POLIMER LINE", "{{legal_name}}"),
        ("20208000305061552001", "{{bank_account}}"),
        ("АК «Хамкорбанк»", "{{bank_name}}"),
        ("АК Хамкорбанк", "{{bank_name}}"),
        ("Хамкорбанк", "{{bank_name}}"),
        ("«03»   02   26г.", "{{claim_date}}"),
        ("№073456н", "№{{base_contract_no}}"),
        ("от 08.08.2023г.", "от {{base_contract_date}}г."),
        ("306338752", "{{inn}}"),
        ("057428н", "{{base_contract_no}}"),
        ("12.05.2023", "{{base_contract_date}}"),
        ("47190", "{{oked}}"),
        ("00083", "{{mfo}}"),
    ]
    n = _xml_replace_paragraphs(src, repls)
    print(f"  scalar substitutions: {n} paragraphs")


def _xml_replace_paragraphs(path: Path, repls: list[tuple[str, str]]) -> int:
    """Replace strings inside word/document.xml at the paragraph level.

    Each <w:p> may have its visible text split across N <w:t> runs because
    Word merges runs by formatting boundaries. The python-docx
    replace-in-paragraph helper concatenates run text then writes back to
    a single run — but only if the FIRST run's joined text contains the
    needle. Some templates split tokens differently and that helper
    misses them. This implementation walks every <w:p>, joins all <w:t>
    children, applies replacements on the joined string, and if changed
    writes the new text into the first <w:t> while clearing the rest.
    """
    import shutil, tempfile
    import xml.etree.ElementTree as ET

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", W)

    work = Path(tempfile.mkdtemp(prefix="reauthor_"))
    try:
        with __import__("zipfile").ZipFile(path, "r") as z:
            z.extractall(work)
        doc_xml = work / "word" / "document.xml"
        tree = ET.parse(doc_xml)
        root = tree.getroot()
        changes = 0
        for p in root.iter(f"{{{W}}}p"):
            text_elems = [t for r in p.findall(f".//{{{W}}}r")
                            for t in r.findall(f"{{{W}}}t")]
            if not text_elems:
                continue
            joined = "".join(t.text or "" for t in text_elems)
            new = joined
            for needle, repl in repls:
                if needle in new:
                    new = new.replace(needle, repl)
            if new == joined:
                continue
            changes += 1
            text_elems[0].text = new
            text_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for t in text_elems[1:]:
                t.text = ""
        tree.write(doc_xml, xml_declaration=True, encoding="UTF-8")
        # Repack
        import zipfile
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
            for base, _, files in __import__("os").walk(work):
                base_p = Path(base)
                for f in files:
                    full = base_p / f
                    z.write(full, full.relative_to(work).as_posix())
        return changes
    finally:
        shutil.rmtree(work, ignore_errors=True)


def reauthor_agreement_uz() -> None:
    src = TEMPLATES / "agreement_uz.docx"
    print(f"\n→ {src.name}")
    doc = Document(src)

    repls = [
        ("(____________________yildagi) -sonli Vositachilik Shartnomasiga (Oferta)",
         "({{base_contract_date}}yildagi) №{{base_contract_no}}-sonli Vositachilik Shartnomasiga (Oferta)"),
        ("(     ) -sonli Qo‘shimcha bitim", "№{{claim_no}}-sonli Qo‘shimcha bitim"),
        ("“30”   04   2025.", "{{claim_date}}"),
    ]
    n = replace_anywhere(doc, repls)
    print(f"  scalar substitutions: {n}")
    doc.save(src)
    print(f"  saved {src}")


def main() -> None:
    if not TEMPLATES.is_dir():
        sys.exit(f"templates dir not found: {TEMPLATES}")

    # Backup originals once. If backup exists, leave alone (idempotent).
    backup_dir = ROOT / "templates" / ".original"
    if not backup_dir.exists():
        backup_dir.mkdir(parents=True, exist_ok=True)
        for f in TEMPLATES.glob("*.docx"):
            shutil.copy2(f, backup_dir / f.name)
        print(f"backed up originals → {backup_dir}")
    else:
        # Restore from backup before re-authoring (so script is idempotent).
        for f in backup_dir.glob("*.docx"):
            shutil.copy2(f, TEMPLATES / f.name)
        print("restored originals from backup")

    reauthor_claim_ru()
    reauthor_claim_uz()
    reauthor_agreement_ru()
    reauthor_agreement_uz()
    print("\nDONE.")


if __name__ == "__main__":
    main()
